import json
from pathlib import Path

from docx import Document

from multi_agentic_rag.config import Settings
from multi_agentic_rag.coverage import plan_requirement_coverage
from multi_agentic_rag.exceptions import IngestionError
from multi_agentic_rag.ingestion import ingest_document
from multi_agentic_rag.ingestion.parser import load_docx_pages
from multi_agentic_rag.models import DocumentStatus
from multi_agentic_rag.storage.sqlite_registry import SQLiteRegistry
from multi_agentic_rag.tasks import handle_task
from multi_agentic_rag.testing import generate_testcases, get_last_test_result, run_testcases


def test_docx_ingestion_chunks_and_extracts_requirement(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    source = _write_docx(tmp_path / "PROJECT_1_BRD_V1.docx")

    pages = load_docx_pages(source)
    result = ingest_document(source, system_name="PROJECT_1", version="v1", settings=settings)
    registry = SQLiteRegistry(settings.sqlite_db_path)
    registry.initialize()

    facts = registry.list_facts(system_name="PROJECT_1", version="v1")
    assert pages[0].extraction_method == "python-docx"
    assert result.chunks_indexed >= 1
    assert any(fact.fact_key == "requirement:REQ-1" for fact in facts)


def test_multiple_documents_can_remain_active_in_same_version(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    first = _write_docx(tmp_path / "PROJECT_1_BRD_V1.docx")
    second = _write_docx(tmp_path / "PROJECT_1_PROTOCOL_V1.docx")

    ingest_document(first, system_name="PROJECT_1", version="v1", settings=settings)
    ingest_document(second, system_name="PROJECT_1", version="v1", settings=settings)

    registry = SQLiteRegistry(settings.sqlite_db_path)
    registry.initialize()
    active_documents = registry.list_documents(
        system_name="PROJECT_1",
        status=DocumentStatus.ACTIVE,
    )

    assert len(active_documents) == 2
    assert {document.source_name for document in active_documents} == {
        "PROJECT_1_BRD_V1.docx",
        "PROJECT_1_PROTOCOL_V1.docx",
    }


def test_ingestion_rejects_filename_version_mismatch(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    source = _write_docx(tmp_path / "PROJECT_1_BRD_V2.docx")

    try:
        ingest_document(source, system_name="PROJECT_1", version="v1", settings=settings)
    except IngestionError as exc:
        assert "Source filename suggests version v2" in str(exc)
    else:
        raise AssertionError("Expected version mismatch to be rejected")


def test_coverage_plan_generates_25_and_reuses_same_scope(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    source = _write_docx(tmp_path / "PROJECT_1_BRD_V1.docx")
    ingest_document(source, system_name="PROJECT_1", version="v1", settings=settings)

    first = plan_requirement_coverage(
        system_name="PROJECT_1",
        version="v1",
        settings=settings,
    )
    second = plan_requirement_coverage(
        system_name="PROJECT_1",
        version="v1",
        settings=settings,
    )

    assert first.supported
    assert first.action == "generated"
    assert len(first.records) == 25
    assert second.action == "reused"
    assert second.run is not None
    assert first.run is not None
    assert second.run.run_id == first.run.run_id


def test_coverage_plan_supports_brd_area_requirement_ids(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    source = _write_brd_style_docx(tmp_path / "PROJECT_1_BRD_STYLE.docx")
    ingest_document(source, system_name="PROJECT_1", version="v1", settings=settings)

    result = plan_requirement_coverage(
        system_name="PROJECT_1",
        version="v1",
        scenario_count=2,
        settings=settings,
    )

    assert result.supported
    assert {record.requirement_id for record in result.records} == {
        "BR-COM-001",
        "BR-SEN-001",
    }
    assert all(record.evidence for record in result.records)


def test_coverage_plan_blocks_when_documents_have_no_requirement_links(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    source = _write_unlinked_docx(tmp_path / "PROJECT_1_NOT_REQUIREMENTS.docx")
    ingest_document(source, system_name="PROJECT_1", version="v1", settings=settings)

    result = plan_requirement_coverage(
        system_name="PROJECT_1",
        version="v1",
        settings=settings,
    )

    assert not result.supported
    assert result.action == "unsupported"
    assert result.message == "No requirement evidence found. No coverage claim can be made."


def test_generated_pytest_file_executes_and_last_result_is_stored(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    source = _write_docx(tmp_path / "PROJECT_1_BRD_V1.docx")
    output_dir = tmp_path / "generated"
    ingest_document(source, system_name="PROJECT_1", version="v1", settings=settings)

    generation = generate_testcases(
        system_name="PROJECT_1",
        version="v1",
        scenario_count=3,
        output_dir=output_dir,
        settings=settings,
    )
    execution = run_testcases(
        system_name="PROJECT_1",
        version="v1",
        scenario_count=3,
        output_dir=output_dir,
        settings=settings,
    )
    last = get_last_test_result(system_name="PROJECT_1", version="v1", settings=settings)

    assert generation.supported
    assert generation.test_file is not None
    generated_file = Path(generation.test_file.file_path)
    tracking_file = Path(generation.test_file.tracking_file_path or "")
    assert generated_file.exists()
    assert generated_file.parent.name == "brd_v1"
    assert generated_file.parent.parent.name == "project_1"
    assert tracking_file.exists()
    assert "class TestProject1V1Automation" in generated_file.read_text(encoding="utf-8")
    assert execution.supported
    assert execution.result is not None
    assert execution.result.status == "blocked"
    assert execution.result.passed == 0
    assert execution.result.skipped == 3
    assert execution.result.failure_category == "PROTOCOL_UNAVAILABLE"
    assert execution.result.dependency_blockers
    tracking = json.loads(tracking_file.read_text(encoding="utf-8"))
    assert tracking["schema_version"] == "test-automation-tracking.v2"
    assert tracking["mode"] == "dependency_aware_generation"
    assert tracking["dependency_audit"]["status"] == "blocked"
    assert tracking["dependency_audit"]["missing_dependencies"]
    assert tracking["protocols"] == ["REST"]
    assert len(tracking["scenarios"]) == 3
    assert len(tracking["selected_scenarios"]) == 3
    assert tracking["run_1"]["status"] == "BLOCKED"
    assert tracking["run_1"]["failure_category"] == "PROTOCOL_UNAVAILABLE"
    assert tracking["db_update_status"] == "test_run_result_record_written"
    assert last.supported
    assert last.result is not None
    assert last.result.result_id == execution.result.result_id


def test_generated_pytest_file_can_execute_in_explicit_mock_mode(tmp_path: Path) -> None:
    settings = _settings(tmp_path, generated_test_execution_mode="mock")
    source = _write_docx(tmp_path / "PROJECT_1_BRD_V1.docx")
    output_dir = tmp_path / "generated"
    ingest_document(source, system_name="PROJECT_1", version="v1", settings=settings)

    execution = run_testcases(
        system_name="PROJECT_1",
        version="v1",
        scenario_count=2,
        output_dir=output_dir,
        settings=settings,
    )

    assert execution.supported
    assert execution.result is not None
    assert execution.result.status == "passed"
    assert execution.result.passed == 2
    tracking = json.loads(Path(execution.tracking_file_path or "").read_text(encoding="utf-8"))
    assert tracking["dependency_audit"]["status"] == "ready"
    assert all(
        scenario["execution_mode"] == "mock"
        for scenario in tracking["selected_scenarios"]
        if scenario["protocols"]
    )


def test_generate_testcases_rewrites_when_requested_count_changes(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    source = _write_docx(tmp_path / "PROJECT_1_BRD_V1.docx")
    output_dir = tmp_path / "generated"
    ingest_document(source, system_name="PROJECT_1", version="v1", settings=settings)

    first = generate_testcases(
        system_name="PROJECT_1",
        version="v1",
        scenario_count=2,
        output_dir=output_dir,
        settings=settings,
    )
    second = generate_testcases(
        system_name="PROJECT_1",
        version="v1",
        scenario_count=3,
        output_dir=output_dir,
        settings=settings,
    )

    assert first.supported
    assert second.supported
    assert second.action == "generated"
    assert second.test_file is not None
    assert len(second.test_file.coverage_ids) == 3
    assert Path(second.test_file.tracking_file_path or "").exists()
    assert _generated_scenario_count(Path(second.test_file.file_path)) == 3


def test_generate_testcases_rewrites_stale_file_for_same_coverage_run(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    source = _write_docx(tmp_path / "PROJECT_1_BRD_V1.docx")
    output_dir = tmp_path / "generated"
    ingest_document(source, system_name="PROJECT_1", version="v1", settings=settings)

    first = generate_testcases(
        system_name="PROJECT_1",
        version="v1",
        scenario_count=3,
        output_dir=output_dir,
        settings=settings,
    )
    assert first.test_file is not None
    file_path = Path(first.test_file.file_path)
    file_path.write_text("COVERAGE_IDS = []\nSCENARIOS = []\n", encoding="utf-8")

    second = generate_testcases(
        system_name="PROJECT_1",
        version="v1",
        scenario_count=3,
        output_dir=output_dir,
        settings=settings,
    )

    assert second.supported
    assert second.action == "generated"
    assert _generated_scenario_count(file_path) == 3


def test_task_router_uses_writer_runner_and_last_result_paths(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    source = _write_docx(tmp_path / "PROJECT_1_BRD_V1.docx")
    output_dir = tmp_path / "generated"
    ingest_document(source, system_name="PROJECT_1", version="v1", settings=settings)

    generate_result = handle_task(
        "Generate testcases for this BRD",
        system_name="PROJECT_1",
        version="v1",
        scenario_count=2,
        output_dir=output_dir,
        settings=settings,
    )
    run_result = handle_task(
        "Run the testcases now",
        system_name="PROJECT_1",
        version="v1",
        scenario_count=2,
        output_dir=output_dir,
        settings=settings,
    )
    last_result = handle_task(
        "Show me the last result",
        system_name="PROJECT_1",
        version="v1",
        settings=settings,
    )

    assert generate_result.intent == "generate_testcases"
    assert generate_result.supported
    assert run_result.intent == "run_testcases"
    assert run_result.supported
    assert last_result.intent == "last_result"
    assert last_result.last_result is not None


def _write_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("REQ-1 The controller shall expose REST GET /api/status.")
    document.add_paragraph("REQ-2 The temperature threshold maximum is 80 C.")
    document.add_paragraph("TEST-1 verifies the status endpoint.")
    document.save(path)
    return path


def _generated_scenario_count(path: Path) -> int:
    return path.read_text(encoding="utf-8").count("'scenario_index':")


def _write_brd_style_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("BR-SEN-001 The controller shall collect sensor data.")
    document.add_paragraph("BR-COM- 001 The controller shall initiate every Modbus transaction.")
    document.save(path)
    return path


def _write_unlinked_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("The controller shall publish readings through MQTT.")
    document.add_paragraph("Temperature threshold maximum is 80 C.")
    document.save(path)
    return path


def _settings(tmp_path: Path, *, generated_test_execution_mode: str = "auto") -> Settings:
    runtime = tmp_path / ".runtime"
    return Settings(
        multi_agentic_rag_home=runtime,
        sqlite_db_path=runtime / "registry.db",
        chroma_path=runtime / "chroma",
        object_store_path=runtime / "objects",
        neo4j_uri=None,
        embedding_provider="hash",
        generated_test_execution_mode=generated_test_execution_mode,
    )
