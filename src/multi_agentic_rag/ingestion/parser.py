"""Document parsing helpers."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from multi_agentic_rag.exceptions import IngestionError
from multi_agentic_rag.utils.paths import resolve_path


@dataclass(frozen=True)
class PageText:
    """Text extracted from one source page."""

    page: int
    text: str
    tables: list[str] = field(default_factory=list)
    extraction_method: str = "pymupdf"


def load_pdf_pages(
    path: str | Path,
    *,
    enable_ocr: bool = False,
    tesseract_cmd: str | None = None,
) -> list[PageText]:
    """Load text and table context from a PDF.

    PyMuPDF is the primary text parser. pdfplumber table extraction is appended
    when available. Tesseract OCR is optional and only attempted for pages with
    no extractable text.
    """

    source = resolve_path(path)
    if not source.exists():
        raise IngestionError(f"Document does not exist: {source}")
    if source.suffix.lower() != ".pdf":
        raise IngestionError("Phase 1 ingestion expects PDF documents.")
    try:
        import fitz
    except Exception as exc:  # pragma: no cover - dependency availability
        raise IngestionError("PyMuPDF is required for PDF ingestion.") from exc

    table_text_by_page = _load_pdf_tables(source)
    pages: list[PageText] = []
    with fitz.open(source) as document:
        for index, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            method = "pymupdf"
            if not text and enable_ocr:
                text = _ocr_page(page, tesseract_cmd=tesseract_cmd)
                if text:
                    method = "tesseract"
            tables = table_text_by_page.get(index, [])
            combined = "\n\n".join(part for part in [text, *tables] if part).strip()
            if combined:
                pages.append(PageText(page=index, text=combined, tables=tables, extraction_method=method))
    if not pages:
        raise IngestionError(f"No extractable text found in {source}")
    return pages


def load_docx_pages(path: str | Path) -> list[PageText]:
    """Load paragraphs and tables from a DOCX document.

    DOCX has no page boundary in the file format, so the whole document is
    represented as page 1 while preserving paragraph and table order as text.
    """

    source = resolve_path(path)
    if not source.exists():
        raise IngestionError(f"Document does not exist: {source}")
    if source.suffix.lower() != ".docx":
        raise IngestionError("DOCX ingestion expects a .docx document.")
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency availability
        raise IngestionError("python-docx is required for DOCX ingestion.") from exc

    document = Document(str(source))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise IngestionError(f"No extractable text found in {source}")
    return [PageText(page=1, text=text, extraction_method="python-docx")]


def load_document_pages(
    path: str | Path,
    *,
    enable_ocr: bool = False,
    tesseract_cmd: str | None = None,
) -> list[PageText]:
    """Load supported source document pages."""

    source = resolve_path(path)
    suffix = source.suffix.lower()
    if suffix == ".pdf":
        return load_pdf_pages(
            source,
            enable_ocr=enable_ocr,
            tesseract_cmd=tesseract_cmd,
        )
    if suffix == ".docx":
        return load_docx_pages(source)
    raise IngestionError("Supported document types are .pdf and .docx.")


def _load_pdf_tables(path: Path) -> dict[int, list[str]]:
    try:
        import pdfplumber
    except Exception:
        return {}

    tables_by_page: dict[int, list[str]] = {}
    try:
        with pdfplumber.open(path) as document:
            for index, page in enumerate(document.pages, start=1):
                page_tables = []
                for table in page.extract_tables() or []:
                    rows = [
                        " | ".join((cell or "").strip() for cell in row).strip()
                        for row in table
                        if row
                    ]
                    rendered = "\n".join(row for row in rows if row)
                    if rendered:
                        page_tables.append(rendered)
                if page_tables:
                    tables_by_page[index] = page_tables
    except Exception:
        return {}
    return tables_by_page


def _ocr_page(page, *, tesseract_cmd: str | None) -> str:
    try:
        from PIL import Image
        import fitz
        import pytesseract
    except Exception:
        return ""

    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    try:
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
        return pytesseract.image_to_string(image).strip()
    except Exception:
        return ""
